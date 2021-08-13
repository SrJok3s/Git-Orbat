#!/usr/bin/env python3

import requests
import re
import colorama
from colorama import Fore, Back, Style, init
import pyfiglet
from docx import Document
from docx.shared import Inches

banner = pyfiglet.figlet_format("Git Orbat") #banner
print(banner)
print(" ")
init()

username=input(Fore.RED+"Send me the target username: "+Style.RESET_ALL)
repo=input(Fore.RED+"Do you have a repo name?: "+Style.RESET_ALL)

# public events scraping
requestGit=requests.get("https://api.github.com/users/{0}/events/public".format(str(username)))	
findEmailField = re.findall(r"([a-zA-Z0-9._-]+@[a-zA-Z0-9._-]+\.[a-zA-Z0-9_-]+)", requestGit.text)
findRepo = re.findall(r"\"html_url\":\"(.*?)\"", requestGit.text)
findAvatar = re.findall(r"\"avatar_url\":\"(.*?)\"", requestGit.text)
findRepoDate= re.findall(r"\"created_at\":\"(.*?)\"", requestGit.text)
finduserLogn=re.findall(r"\"login\":\"(.*?)\"", requestGit.text)
reponame=[]
create_date_repo=[]
avatar_url=[]
y=0
# User info all scrape
userall=requests.get(f"https://api.github.com/users/{username}")
findType = re.findall(r"\"type\":\"(.*?)\"", userall.text)
findName = re.findall(r"\"login\":\"(.*?)\"", userall.text)
findLocation = re.findall(r"\"location\":\"(.*?)\"", userall.text)
findbio = re.findall(r"\"bio\":\"(.*?)\"", userall.text)
findpubrepos = re.findall(r"\"public_repos\":\"(.*?)\"", userall.text)
findgists = re.findall(r"\"public_gists\":\"(.*?)\"", userall.text)
findCreation = re.findall(r"\"created_at\":\"(.*?)\"", userall.text)
findUpdated = re.findall(r"\"updated_at\":\"(.*?)\"", userall.text)
findCompany = re.findall(r"\"company\":\"(.*?)\"", userall.text)
finduserid = re.findall(r"\"id\":\"(.*?)\"", userall.text)

while (y<1):
	avatar_url.append(findAvatar[y])
	y=y+1 # Print [0] unicamente
reponame.append(findRepo) # Pueden localizarse mas de uno, array + printeo todos
create_date_repo.append(findRepoDate)

requestComits=requests.get(f"https://api.github.com/repos/{username}/{repo}/commits")
findEmailField_2 = re.findall(r"([a-zA-Z0-9._-]+@[a-zA-Z0-9._-]+\.[a-zA-Z0-9_-]+)", requestComits.text)

# Deleting repeated emails
findEmailField = list(dict.fromkeys(findEmailField))
findEmailField_2 = list(dict.fromkeys(findEmailField_2))


print(" ")
print(Fore.RED+" User information found: "+Style.RESET_ALL)
y=0
print(" ")

while (y < 1):
	print(Fore.GREEN+" Avatar url -> "+str(avatar_url[y])+Style.RESET_ALL)
	y=y+1
print(Fore.GREEN+" User ID -> "+str(finduserid))
print(" User Type -> "+str(findType))
print(" Name -> "+str(findName))
print(" Company -> "+str(findCompany))
print(" Location -> "+str(findLocation))
print(" Bio -> "+str(findbio))
print(" Public Repos -> "+str(findpubrepos))
print(" Public Gists -> "+str(findgists))
print(" Created at -> "+str(findCreation))
print(" Last Updated -> " +str(findUpdated)+Style.RESET_ALL)
print(" ")
print(" Repo information found -> ")
print(" ")
repoinfo=requests.get(f"https://api.github.com/users/{username}/repos")
findRepoName = re.findall(r"\"full_name\":\"(.*?)\"", repoinfo.text)
findCreationRepo = re.findall(r"\"created_at\":\"(.*?)\"", repoinfo.text)
reponombre=[]
repoCreation=[]
for x in range(len(findRepoName)):
	reponombre.append(findRepoName[x])
	repoCreation.append(findCreationRepo[x])
	x=x+1
total= len(reponombre)
x=0
for fine in reponombre:
	print(Fore.GREEN+str(reponombre[x])+" | "+str(repoCreation[x])+Style.RESET_ALL)
	print(" ")
	x=x+1
# Printing emails
print(Fore.RED+"Emails found: "+Style.RESET_ALL)
print(" ")
y = 0
for findEmail in findEmailField:
	print("-> "+Fore.GREEN+str(findEmailField[y])+Style.RESET_ALL)
	y=y+1
y = 0
for findE in findEmailField_2:
	print("-> "+Fore.GREEN+str(findEmailField_2[y])+Style.RESET_ALL)
	y=y+1

requestFollowers=requests.get("https://api.github.com/users/{0}/followers".format(str(username)))
findLogin = re.findall(r"\"login\":\"(.*?)\"", requestFollowers.text)
followers=[]
print(" ")
followers.append(findLogin)
print(" ")
print(Fore.RED+"Extracting followers: "+Style.RESET_ALL)
y=0
for x in findLogin:
	print(Fore.GREEN+"-> "+findLogin[y]+Style.RESET_ALL)
	y=y+1

requestFollowing=requests.get("https://api.github.com/users/{0}/following".format(str(username)))
findLogin_f = re.findall(r"\"login\":\"(.*?)\"", requestFollowing.text)
following=[]
following.append(findLogin_f)
print(" ")
print(Fore.RED+"Extracting following: "+Style.RESET_ALL)
y=0
for x in findLogin_f:
	print(Fore.GREEN+"-> "+findLogin_f[y]+Style.RESET_ALL)
	y=y+1

# Empieza a generar el documento
print("---------------------------------------------------")
print("------------ Generating Document ------------------")
document = Document()
document.add_heading('Análisis perfil '+str(username), 0)
document.add_paragraph('A continuación se analiza el perfil indicado y se muestran los resultados e informació relevante detectada: ')
document.add_heading('Información básica', level=1)
document.add_paragraph('Avatar url: '+str(avatar_url[0]))
document.add_paragraph('User ID: '+str(finduserid))
document.add_paragraph('Bio: '+str(findbio))
document.add_paragraph('Profile Created at: '+str(findCreation))
document.add_paragraph('User Type:  '+str(findType))
document.add_paragraph('Company: '+str(findCompany))
document.add_paragraph('Location: '+str(findLocation))
document.add_heading('Correos detectados ', level=1)
y=0
for findEmail in findEmailField:
	document.add_paragraph('-> '+str(findEmailField[y]))
	y=y+1
y = 0
for findE in findEmailField_2:
	document.add_paragraph('-> '+str(findEmailField_2[y]))
	y=y+1
document.add_heading('Repo info found ', level=1)
for y in range(len(reponombre)):
	document.add_paragraph('Repo name: '+str(reponombre[y])+" |  Creation Date: "+str(findCreationRepo[y]))
document.add_heading("Following: ")
for y in range(len(following)):
	document.add_paragraph("· "+str(following[y]))

document.add_heading("Followers: ")
for y in range(len(followers)):
	document.add_paragraph("· "+str(followers[y]))
document.add_page_break()
document.save('report_github.docx')
